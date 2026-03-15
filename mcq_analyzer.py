"""
╔══════════════════════════════════════════════════════════════╗
║           MCQ ANALYZER  —  Professional Dashboard           ║
║                                                              ║
║  Supported Structure:                                        ║
║    Q 1. Statement of question                                ║
║    (A) option one   (B) option two                           ║
║    (C) option three (D) option four                          ║
║    Answer: A                                                 ║
║    Explanations: text...                                     ║
║                                                              ║
║  Run:  python mcq_analyzer.py                                ║
╚══════════════════════════════════════════════════════════════╝
"""

import re, sys, os, math, time, threading, json, webbrowser
from collections import defaultdict
try:
    import urllib.request as _urllib_req
except ImportError:
    _urllib_req = None

# ══════════════════════════════════════════════════════════════
#  VERSION & GITHUB UPDATE CONFIG
#  ─────────────────────────────────────────────────────────────
#  SETUP KARNE KA TARIKA:
#  1. GitHub pe repo banao: github.com/YourName/mcq-analyzer
#  2. Repo mein ek file banao:  version.json  with content:
#     {
#       "version": "2.1",
#       "release_notes": "New features added",
#       "download_url": "https://github.com/YourName/mcq-analyzer/releases/latest"
#     }
#  3. Neeche GITHUB_USER aur GITHUB_REPO apna likhو
#  4. Har naye update pe version.json mein version badlo
# ══════════════════════════════════════════════════════════════

CURRENT_VERSION = "2.0"
GITHUB_USER     = "veershehryal9t9-arch"
GITHUB_REPO     = "software"
VERSION_URL     = (f"https://raw.githubusercontent.com/"
                   f"{GITHUB_USER}/{GITHUB_REPO}/main/version.json")
DOWNLOAD_URL    = (f"https://github.com/{GITHUB_USER}/"
                   f"{GITHUB_REPO}/releases/latest")

# ══════════════════════════════════════════════════════════════
#  LICENSE CONFIG
#  ─────────────────────────────────────────────────────────────
#  Apna server URL yahan daalo jab deploy ho jaye
#  e.g. https://mcq-license.up.railway.app
# ══════════════════════════════════════════════════════════════

LICENSE_SERVER  = "https://your-server.up.railway.app"
LICENSE_FILE    = os.path.join(os.path.expanduser("~"),
                               ".mcq_analyzer_license.json")

# Trial mein yeh features available hain
TRIAL_FEATURES = {
    "max_questions" : 20,       # sirf 20 questions analyze
    "no_report_save": True,     # report save nahi hogi
    "no_layout"     : True,     # layout check nahi hoga
    "watermark"     : True,     # TRIAL watermark
}

# ── Hardware ID (machine fingerprint) ────────────────────────
def get_hwid():
    """Current PC ka unique hardware ID banata hai."""
    import hashlib, platform, uuid
    raw = (platform.node() +
           platform.processor() +
           str(uuid.getnode()))
    return hashlib.sha256(raw.encode()).hexdigest()[:32]

# ── License local cache save/load ────────────────────────────
def save_license_cache(data):
    try:
        with open(LICENSE_FILE, "w") as f:
            json.dump(data, f)
    except Exception:
        pass

def load_license_cache():
    try:
        if os.path.exists(LICENSE_FILE):
            with open(LICENSE_FILE, "r") as f:
                return json.load(f)
    except Exception:
        pass
    return None

# ── Validate license with server ─────────────────────────────
def validate_license_online(key):
    """
    Server se license validate karta hai.
    Returns dict: {licensed, status, message, plan, name}
    """
    if _urllib_req is None:
        return {"licensed": False, "status": "error",
                "message": "Network not available"}
    try:
        hwid    = get_hwid()
        payload = json.dumps({
            "license_key": key,
            "hwid":        hwid
        }).encode()
        req = _urllib_req.Request(
            f"{LICENSE_SERVER}/api/validate",
            data=payload,
            headers={"Content-Type": "application/json",
                     "User-Agent":   "MCQ-Analyzer/2.0"})
        with _urllib_req.urlopen(req, timeout=8) as resp:
            return json.loads(resp.read().decode())
    except Exception as e:
        return {"licensed": False, "status": "error",
                "message": str(e)}

# ── Check license status ─────────────────────────────────────
def check_license():
    """
    Returns: (is_licensed:bool, is_trial:bool, info:dict)
    """
    cache = load_license_cache()

    if cache and cache.get("licensed"):
        # Cached license hai — online recheck karo
        result = validate_license_online(cache.get("key", ""))
        if result.get("licensed"):
            save_license_cache({**result, "key": cache["key"]})
            return True, False, result
        else:
            # Cache clear karo agar revoked
            if result.get("status") == "revoked":
                save_license_cache({})
            # Offline fallback — agar server down ho
            if result.get("status") == "error":
                return True, False, cache   # trust cache
            return False, True, result

    # Koi license nahi — trial mode
    return False, True, {"status": "trial",
                         "message": "Trial Mode",
                         "plan":    "trial"}

# ══════════════════════════════════════════════════════════════
#  LICENSE ACTIVATION WINDOW
# ══════════════════════════════════════════════════════════════

def show_license_window(parent_win=None, on_success=None):
    """Professional license activation window."""
    import tkinter as tk

    W, H = 560, 420
    lwin = tk.Toplevel(parent_win) if parent_win else tk.Tk()
    lwin.overrideredirect(True)
    lwin.configure(bg="#0d1117")
    sw = lwin.winfo_screenwidth(); sh = lwin.winfo_screenheight()
    lwin.geometry(f"{W}x{H}+{(sw-W)//2}+{(sh-H)//2}")
    lwin.attributes("-topmost", True)
    lwin.grab_set()

    # borders
    tk.Frame(lwin, bg="#d29922", height=3).place(x=0, y=0, width=W)
    tk.Frame(lwin, bg="#238636", height=3).place(x=0, y=H-3, width=W)
    tk.Frame(lwin, bg="#30363d", width=1).place(x=0,   y=0, height=H)
    tk.Frame(lwin, bg="#30363d", width=1).place(x=W-1, y=0, height=H)

    # title bar
    tb = tk.Frame(lwin, bg="#161b22", height=36)
    tb.place(x=1, y=3, width=W-2)
    tk.Label(tb, text="  MCQ Analyzer  —  License Activation",
             font=("Segoe UI", 10, "bold"),
             bg="#161b22", fg="#d29922").pack(side="left", padx=6, pady=8)
    tk.Button(tb, text="✕", command=lwin.destroy,
              bg="#161b22", fg="#f85149",
              font=("Segoe UI", 10, "bold"),
              relief="flat", bd=0, padx=10,
              activebackground="#2d1c1c",
              cursor="hand2").pack(side="right")

    # lock icon
    lk = tk.Canvas(lwin, width=60, height=60,
                   bg="#0d1117", highlightthickness=0)
    lk.place(relx=0.5, rely=0.20, anchor="center")
    lk.create_oval(10,25,50,58, fill="#161b22", outline="#d29922", width=2)
    lk.create_arc(15,8,45,35, start=0, extent=180,
                  style="arc", outline="#d29922", width=3)
    lk.create_rectangle(26,35,34,48, fill="#d29922", outline="")

    tk.Label(lwin, text="Activate Your License",
             font=("Segoe UI", 14, "bold"),
             bg="#0d1117", fg="#e6edf3").place(
             relx=0.5, rely=0.38, anchor="center")

    tk.Label(lwin,
             text="Enter your license key to unlock all features.",
             font=("Segoe UI", 9),
             bg="#0d1117", fg="#8b949e").place(
             relx=0.5, rely=0.47, anchor="center")

    # key entry box
    key_var = tk.StringVar()
    key_entry = tk.Entry(
        lwin,
        textvariable=key_var,
        font=("Courier New", 12, "bold"),
        bg="#161b22", fg="#d29922",
        insertbackground="#d29922",
        relief="flat", bd=0,
        justify="center",
        width=28
    )
    key_entry.place(relx=0.5, rely=0.58, anchor="center",
                    width=380, height=38)
    # entry border frame
    tk.Frame(lwin, bg="#d29922", height=2).place(
        relx=0.5, rely=0.625, anchor="center", width=380)

    key_entry.insert(0, "MCQA-XXXX-XXXX-XXXX-XXXX")
    key_entry.bind("<FocusIn>", lambda e: (
        key_entry.delete(0, "end"),
        key_entry.config(fg="#d29922")
    ))

    # status label
    status_lbl = tk.Label(lwin, text="",
                          font=("Segoe UI", 9, "bold"),
                          bg="#0d1117", fg="#8b949e")
    status_lbl.place(relx=0.5, rely=0.70, anchor="center")

    # progress bar
    pb = tk.Canvas(lwin, width=380, height=4,
                   bg="#21262d", highlightthickness=0)
    pb.place(relx=0.5, rely=0.76, anchor="center")
    pb_bar = pb.create_rectangle(0, 0, 0, 4,
                                 fill="#d29922", outline="")

    # Activate button
    def _activate():
        key = key_var.get().strip().upper()
        if len(key) < 10:
            status_lbl.config(text="Please enter a valid license key.",
                              fg="#f85149")
            return

        act_btn.config(state="disabled", text="  Validating...")
        status_lbl.config(text="Connecting to license server...",
                          fg="#8b949e")

        # animated bar
        _apos = [0.0]
        _animating = [True]
        def _anim():
            if not _animating[0]: return
            _apos[0] = (_apos[0]+0.04) % 1.2
            t = _apos[0]
            pb.coords(pb_bar,
                      int(380*max(0,t-0.3)), 0,
                      int(380*min(1,t)), 4)
            lwin.after(16, _anim)
        lwin.after(20, _anim)

        def _check():
            result = validate_license_online(key)
            _animating[0] = False

            def _show(r=result):
                pb.coords(pb_bar, 0, 0, 380, 4)
                if r.get("licensed"):
                    pb.itemconfig(pb_bar, fill="#3fb950")
                    status_lbl.config(
                        text=f"  Activated! Welcome, {r.get('name','User')}!",
                        fg="#3fb950")
                    save_license_cache({**r, "key": key})
                    act_btn.config(text="  Activated!", bg="#3fb950")
                    lwin.after(1200, lambda: (
                        lwin.destroy(),
                        on_success() if on_success else None
                    ))
                else:
                    pb.itemconfig(pb_bar, fill="#f85149")
                    status_lbl.config(
                        text=f"  {r.get('message','Invalid key.')}",
                        fg="#f85149")
                    act_btn.config(
                        state="normal", text="  Activate",
                        bg="#d29922")
            lwin.after(0, _show)

        threading.Thread(target=_check, daemon=True).start()

    act_btn = tk.Button(
        lwin, text="  Activate",
        command=_activate,
        bg="#d29922", fg="#0d1117",
        font=("Segoe UI", 11, "bold"),
        relief="flat", bd=0,
        padx=28, pady=9,
        cursor="hand2",
        activebackground="#e3a11a"
    )
    act_btn.place(relx=0.38, rely=0.87, anchor="center")
    act_btn.bind("<Enter>", lambda e: act_btn.config(bg="#e3a11a"))
    act_btn.bind("<Leave>", lambda e: act_btn.config(bg="#d29922"))

    # Continue trial button
    def _trial():
        lwin.destroy()

    tk.Button(
        lwin, text="  Continue Trial",
        command=_trial,
        bg="#21262d", fg="#8b949e",
        font=("Segoe UI", 10),
        relief="flat", bd=0,
        padx=18, pady=9,
        cursor="hand2",
        activebackground="#30363d"
    ).place(relx=0.68, rely=0.87, anchor="center")

    # trial info
    tk.Label(lwin,
             text=f"Trial: max {TRIAL_FEATURES['max_questions']} questions  |  No report save  |  No layout check",
             font=("Segoe UI", 7),
             bg="#0d1117", fg="#444").place(
             relx=0.5, rely=0.96, anchor="center")

    key_entry.focus()
    if parent_win is None:
        lwin.mainloop()


def check_for_update():
    """GitHub se version.json fetch karke compare karta hai."""
    if _urllib_req is None:
        return {"has_update": False, "latest": "?",
                "current": CURRENT_VERSION, "notes": "",
                "download_url": DOWNLOAD_URL,
                "error": "urllib not available"}
    try:
        req = _urllib_req.Request(
            VERSION_URL,
            headers={"User-Agent": "MCQ-Analyzer-Updater/2.0"})
        with _urllib_req.urlopen(req, timeout=8) as resp:
            data = json.loads(resp.read().decode())
        latest = data.get("version", "0.0").strip()
        notes  = data.get("release_notes", "New update available.")
        dl_url = data.get("download_url", DOWNLOAD_URL)

        def ver_tuple(v):
            try:
                return tuple(int(x) for x in str(v).split("."))
            except Exception:
                return (0,)

        has_update = ver_tuple(latest) > ver_tuple(CURRENT_VERSION)
        return {"has_update": has_update, "latest": latest,
                "current": CURRENT_VERSION, "notes": notes,
                "download_url": dl_url, "error": None}
    except Exception as e:
        return {"has_update": False, "latest": "?",
                "current": CURRENT_VERSION, "notes": "",
                "download_url": DOWNLOAD_URL,
                "error": str(e)}

# ══════════════════════════════════════════════════════════════
#  1. FILE READING
# ══════════════════════════════════════════════════════════════

def read_docx(path):
    try:
        import docx
    except ImportError:
        _fatal("python-docx not found.\n\nPlease run:\n  pip install python-docx")
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs)

def read_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def load_file(path):
    ext = os.path.splitext(path)[1].lower()
    if ext in (".docx", ".doc"):
        return read_docx(path)
    elif ext == ".txt":
        return read_txt(path)
    else:
        _fatal(f"Unsupported format: {ext}\n\nUse .docx or .txt files.")

def _fatal(msg):
    try:
        import tkinter as tk
        from tkinter import messagebox
        r = tk.Tk(); r.withdraw()
        messagebox.showerror("MCQ Analyzer — Error", msg)
        r.destroy()
    except Exception:
        print("[ERROR]", msg)
    sys.exit(1)

# ══════════════════════════════════════════════════════════════
#  2. FLEXIBLE PARSER
# ══════════════════════════════════════════════════════════════

# Roman numeral → integer
_ROMAN = {"i":1,"ii":2,"iii":3,"iv":4,"v":5,"vi":6,"vii":7,"viii":8,
          "ix":9,"x":10,"xi":11,"xii":12,"xiii":13,"xiv":14,"xv":15,
          "xvi":16,"xvii":17,"xviii":18,"xix":19,"xx":20,
          "xxi":21,"xxii":22,"xxiii":23,"xxiv":24,"xxv":25,
          "xxvi":26,"xxvii":27,"xxviii":28,"xxix":29,"xxx":30,
          "xxxi":31,"xxxii":32,"xxxiii":33,"xxxiv":34,"xxxv":35,
          "xl":40,"l":50,"lx":60,"lxx":70,"lxxx":80,"xc":90,"c":100}

def _roman_to_int(s):
    return _ROMAN.get(s.lower().strip(), None)

# Accepted option letters (A-Z single char)
VALID_OPTS = set("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
STANDARD_OPTS = {"A","B","C","D"}

# Accepted answer letters
VALID_ANS = set("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
STANDARD_ANS = {"A","B","C","D"}

def _parse_qnum(line):
    """
    Tries all question number formats. Returns (int_num, rest_of_line) or None.
    Formats:
      Q 1.  Q1.  Q 1)  Q1)  Q.1  q1
      1.    1)   1:
      Question 1.  Question 1)
      i.  ii.  iii)  (roman)
    """
    # Q / q prefix with optional space/dot
    m = re.match(r'^[Qq](?:uestion)?\s*\.?\s*(\d+)[.):\s](.*)$', line)
    if m:
        return int(m.group(1)), m.group(2).strip()

    # plain digit
    m = re.match(r'^(\d+)[.):\s]\s*(.+)$', line)
    if m:
        return int(m.group(1)), m.group(2).strip()

    # roman numeral  i.  ii)  iii.
    m = re.match(r'^([IiVvXxLlCc]{1,8})[.):\s]\s*(.+)$', line)
    if m:
        n = _roman_to_int(m.group(1))
        if n:
            return n, m.group(2).strip()

    return None

def _parse_option(line):
    """
    Tries all option formats. Returns (letter_upper, text) or None.
    Formats:
      (A) text    (a) text
      A) text     a) text
      A. text     a. text
      A- text
      A  text   (letter + 2+ spaces)
    """
    # (X) or (x)
    m = re.match(r'^\(([A-Za-z])\)\s*(.+)$', line)
    if m:
        return m.group(1).upper(), m.group(2).strip()

    # X) or X. or X-
    m = re.match(r'^([A-Za-z])[.)–\-]\s*(.+)$', line)
    if m:
        return m.group(1).upper(), m.group(2).strip()

    # X   text  (letter + 2+ spaces)
    m = re.match(r'^([A-Za-z])\s{2,}(.+)$', line)
    if m:
        return m.group(1).upper(), m.group(2).strip()

    return None

def _parse_answer(line):
    """
    Parses Answer line. Returns letter (upper) or None.
    Detects non-standard letters too.
    """
    m = re.match(r'^(?:Answer|Ans|Key)\s*[:=.]\s*([A-Za-z])\s*$', line, re.IGNORECASE)
    if m:
        return m.group(1).upper()
    # Answer: A. text  (letter followed by other stuff)
    m = re.match(r'^(?:Answer|Ans|Key)\s*[:=.]\s*([A-Za-z])[^a-zA-Z]', line, re.IGNORECASE)
    if m:
        return m.group(1).upper()
    return None

def parse_mcqs(text):
    lines     = [l.strip() for l in text.splitlines()]
    questions = []
    current   = None
    EXP_PAT   = re.compile(r'^Explanation[s]?\s*[:=]\s*(.*)', re.IGNORECASE)

    for line in lines:
        if not line:
            continue

        # ── Try question number ───────────────────────────────
        qr = _parse_qnum(line)
        if qr:
            if current:
                questions.append(current)
            num, stmt = qr
            current = {"num": num, "statement": stmt,
                       "options": {}, "answer": None, "explanation": None,
                       "invalid_opts": [],   # options with non-ABCD letter
                       "invalid_ans":  None} # answer with non-ABCD letter
            continue

        if current is None:
            continue

        # ── Try answer ────────────────────────────────────────
        ans = _parse_answer(line)
        if ans:
            current["answer"] = ans
            if ans not in STANDARD_ANS:
                current["invalid_ans"] = ans
            continue

        # ── Try explanation ───────────────────────────────────
        me = EXP_PAT.match(line)
        if me:
            current["explanation"] = me.group(1).strip() or "(present)"
            continue

        # ── Try option ────────────────────────────────────────
        opt = _parse_option(line)
        if opt:
            letter, text = opt
            current["options"][letter] = text
            if letter not in STANDARD_OPTS:
                current["invalid_opts"].append(letter)
            continue

        # ── Statement continuation ────────────────────────────
        if current["statement"] and not current["options"]:
            current["statement"] += " " + line

    if current:
        questions.append(current)
    return questions

# ══════════════════════════════════════════════════════════════
#  2b. LAYOUT ANALYSER  (for .docx files)
# ══════════════════════════════════════════════════════════════

def analyze_layout(filepath):
    """
    Analyses .docx for:
      - Line spacing per paragraph
      - Page margins (top/bottom/left/right in cm)
      - Missing page numbers (checks for page number fields)
      - Detected page count
    Returns a dict.  For .txt files returns None.
    """
    ext = os.path.splitext(filepath)[1].lower()
    if ext not in (".docx", ".doc"):
        return None
    try:
        import docx
        from docx.oxml.ns import qn
        doc = docx.Document(filepath)

        # ── margins ──────────────────────────────────────────
        sec = doc.sections[0]
        def emu_to_cm(emu): return round(emu / 914400 * 2.54, 2) if emu else 0
        margins = {
            "top":    emu_to_cm(sec.top_margin),
            "bottom": emu_to_cm(sec.bottom_margin),
            "left":   emu_to_cm(sec.left_margin),
            "right":  emu_to_cm(sec.right_margin),
        }

        # ── line spacing ─────────────────────────────────────
        spacings = []
        for para in doc.paragraphs:
            pf = para.paragraph_format
            if pf.line_spacing is not None:
                try:
                    val = float(pf.line_spacing)
                    if val > 0:
                        spacings.append(round(val, 2))
                except Exception:
                    pass
        if spacings:
            from collections import Counter
            most_common = Counter(spacings).most_common(1)[0][0]
            line_spacing = most_common
        else:
            line_spacing = "Default (single)"

        # ── page numbers ─────────────────────────────────────
        has_page_num = False
        for el in doc.element.iter():
            if el.tag == qn("w:fldChar") or el.tag == qn("w:instrText"):
                txt = el.text or ""
                if "PAGE" in txt.upper():
                    has_page_num = True
                    break
        # also check headers/footers
        for sec2 in doc.sections:
            for hdr in [sec2.header, sec2.footer]:
                if hdr:
                    for para in hdr.paragraphs:
                        for run in para.runs:
                            if "PAGE" in run.text.upper():
                                has_page_num = True
            if has_page_num:
                break

        # ── paragraph count as proxy for pages ───────────────
        para_count = len(doc.paragraphs)
        est_pages  = max(1, para_count // 25)

        return {
            "margins":      margins,
            "line_spacing": line_spacing,
            "has_page_num": has_page_num,
            "para_count":   para_count,
            "est_pages":    est_pages,
        }
    except Exception as e:
        return {"error": str(e)}

# ══════════════════════════════════════════════════════════════
#  3. ANALYSIS
# ══════════════════════════════════════════════════════════════

def normalize(s):
    return re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', s.lower())).strip()

def analyze(questions):
    r = {"total": len(questions), "missing_options": [], "missing_answer": [],
         "missing_expln": [], "missing_qnums": [], "duplicates": [],
         "invalid_opts": [],   # (qnum, [letters])
         "invalid_ans":  [],   # (qnum, letter)
         "questions": questions}

    for q in questions:
        n = q["num"]
        if len(q["options"]) < 4:
            r["missing_options"].append((n, len(q["options"])))
        if not q["answer"]:
            r["missing_answer"].append(n)
        if not q["explanation"]:
            r["missing_expln"].append(n)
        # non-standard option letters
        bad_opts = q.get("invalid_opts", [])
        if bad_opts:
            r["invalid_opts"].append((n, bad_opts))
        # non-standard answer letter
        bad_ans = q.get("invalid_ans")
        if bad_ans:
            r["invalid_ans"].append((n, bad_ans))

    nums = sorted(q["num"] for q in questions)
    if nums:
        r["missing_qnums"] = sorted(set(range(nums[0], nums[-1]+1)) - set(nums))

    seen = defaultdict(list)
    for q in questions:
        k = normalize(q["statement"])
        if len(k) > 8:
            seen[k].append(q["num"])
    for k, qn in seen.items():
        if len(qn) > 1:
            r["duplicates"].append({"qnums": sorted(qn),
                                    "preview": k[:90]+("..." if len(k) > 90 else "")})
    return r

# ══════════════════════════════════════════════════════════════
#  4. REPORT TEXT
# ══════════════════════════════════════════════════════════════

SEP  = "=" * 64
SEP2 = "-" * 64

def build_report(r, filepath, layout=None):
    lines = []; a = lines.append
    q   = r["questions"];  tot = r["total"]
    mo  = r["missing_options"]; ma = r["missing_answer"]
    me  = r["missing_expln"];   mq = r["missing_qnums"]
    dg  = r["duplicates"]
    ivo = r.get("invalid_opts", [])
    iva = r.get("invalid_ans",  [])
    nums_present = sorted(x["num"] for x in q)
    q_range = f"Q{nums_present[0]} to Q{nums_present[-1]}" if nums_present else "N/A"

    a(SEP); a("                 MCQ ANALYSIS REPORT"); a(SEP)
    a(f"  File  : {os.path.basename(filepath)}"); a(SEP); a("")
    a("SUMMARY"); a(SEP2)
    a(f"  Total questions detected      : {tot}")
    a(f"  Question range                : {q_range}")
    a(f"  Questions with all 4 options  : {tot - len(mo)}")
    a(f"  Questions missing options     : {len(mo)}")
    a(f"  Questions missing Answer key  : {len(ma)}")
    a(f"  Questions missing Explanation : {len(me)}")
    a(f"  Missing Q numbers in range    : {len(mq)}")
    a(f"  Duplicate statement groups    : {len(dg)}")
    a(f"  Invalid option letters (non-ABCD): {len(ivo)}")
    a(f"  Invalid answer letters (non-ABCD): {len(iva)}"); a("")

    a(SEP2); a("QUESTIONS WITH MISSING OPTIONS"); a(SEP2)
    if not mo:
        a("  [OK] All questions have 4 options.")
    else:
        for qnum, found in mo:
            s = next((x["statement"][:55] for x in q if x["num"] == qnum), "")
            a(f"  Q{qnum:>4}  ->  only {found} option(s)  |  {s}...")
    a("")

    a(SEP2); a("QUESTIONS WITH MISSING ANSWER KEY"); a(SEP2)
    if not ma:
        a("  [OK] All questions have Answer key.")
    else:
        for qnum in ma:
            s = next((x["statement"][:55] for x in q if x["num"] == qnum), "")
            a(f"  Q{qnum:>4}  ->  Answer: MISSING  |  {s}...")
    a("")

    a(SEP2); a("QUESTIONS WITH MISSING EXPLANATION"); a(SEP2)
    if not me:
        a("  [OK] All questions have Explanation.")
    else:
        for qnum in me:
            s = next((x["statement"][:55] for x in q if x["num"] == qnum), "")
            a(f"  Q{qnum:>4}  ->  Explanation: MISSING  |  {s}...")
    a("")

    a(SEP2); a("MISSING QUESTION NUMBERS IN RANGE"); a(SEP2)
    if not mq:
        a("  [OK] No missing numbers in range.")
    else:
        ranges, start, end = [], mq[0], mq[0]
        for n in mq[1:]:
            if n == end+1: end = n
            else:
                ranges.append(f"Q{start}-Q{end}" if start != end else f"Q{start}")
                start = end = n
        ranges.append(f"Q{start}-Q{end}" if start != end else f"Q{start}")
        a(f"  Missing: {', '.join(ranges)}")
        a(f"  Total  : {len(mq)}")
    a("")

    a(SEP2); a("DUPLICATE / REPEATED QUESTIONS"); a(SEP2)
    if not dg:
        a("  [OK] No duplicates found.")
    else:
        for i, grp in enumerate(dg, 1):
            a(f"  Group {i}: {', '.join(f'Q{n}' for n in grp['qnums'])}")
            a(f"    \"{grp['preview']}\""); a("")
    a("")

    a(SEP2); a("INVALID OPTION LETTERS  (non A/B/C/D)"); a(SEP2)
    if not ivo:
        a("  [OK] All option letters are standard (A/B/C/D).")
    else:
        for qnum, letters in ivo:
            s = next((x["statement"][:50] for x in q if x["num"] == qnum), "")
            a(f"  Q{qnum:>4}  ->  Non-standard options: {', '.join(letters)}  |  {s}...")
    a("")

    a(SEP2); a("INVALID ANSWER LETTERS  (non A/B/C/D)"); a(SEP2)
    if not iva:
        a("  [OK] All answer keys are standard (A/B/C/D).")
    else:
        for qnum, letter in iva:
            s = next((x["statement"][:50] for x in q if x["num"] == qnum), "")
            a(f"  Q{qnum:>4}  ->  Answer '{letter}' is non-standard  |  {s}...")
    a("")

    if layout and not layout.get("error"):
        a(SEP2); a("DOCUMENT LAYOUT ANALYSIS"); a(SEP2)
        mg = layout.get("margins", {})
        a(f"  Margin Top    : {mg.get('top','?')} cm")
        a(f"  Margin Bottom : {mg.get('bottom','?')} cm")
        a(f"  Margin Left   : {mg.get('left','?')} cm")
        a(f"  Margin Right  : {mg.get('right','?')} cm")
        a(f"  Line Spacing  : {layout.get('line_spacing','?')}")
        a(f"  Page Numbers  : {'Present' if layout.get('has_page_num') else 'MISSING'}")
        a(f"  Est. Pages    : {layout.get('est_pages','?')}")
        a(f"  Paragraphs    : {layout.get('para_count','?')}")
        a("")

    a(SEP); a("  END OF REPORT"); a(SEP)
    return "\n".join(lines)

# ══════════════════════════════════════════════════════════════
#  5.  COLOUR PALETTE
# ══════════════════════════════════════════════════════════════

C = {
    "bg"          : "#0d1117",
    "bg2"         : "#161b22",
    "bg3"         : "#21262d",
    "border"      : "#30363d",
    "accent"      : "#58a6ff",
    "accent2"     : "#3fb950",
    "text"        : "#e6edf3",
    "text2"       : "#8b949e",
    "red"         : "#f85149",
    "orange"      : "#d29922",
    "purple"      : "#a371f7",
    "teal"        : "#39d353",
    "card_blue"   : "#1c2d3f",
    "card_red"    : "#2d1c1c",
    "card_orange" : "#2d2310",
    "card_purple" : "#231c2d",
    "card_green"  : "#1c2d1c",
    "btn_primary" : "#238636",
    "btn_prim_hov": "#2ea043",
    "btn_second"  : "#21262d",
    "btn_sec_hov" : "#30363d",
}

# ══════════════════════════════════════════════════════════════
#  6.  GUI HELPERS
# ══════════════════════════════════════════════════════════════

class AnimatedCounter:
    def __init__(self, label, target, color, duration=900):
        self.label    = label
        self.target   = target
        self.color    = color
        self.duration = duration
        self.start    = time.time()
        self._tick()

    def _tick(self):
        elapsed = (time.time() - self.start) * 1000
        if elapsed >= self.duration:
            self.label.config(text=str(self.target), fg=self.color)
            return
        t = elapsed / self.duration
        t = 1 - (1 - t) ** 3
        self.label.config(text=str(int(self.target * t)), fg=self.color)
        self.label.after(16, self._tick)


class AnimatedBar:
    def __init__(self, canvas, x, y, w, h, pct, color, duration=900):
        self.canvas   = canvas
        self.x, self.y, self.w, self.h = x, y, w, h
        self.pct      = pct
        self.color    = color
        self.duration = duration
        self.start    = time.time()
        canvas.create_rectangle(x, y, x+w, y+h, fill=C["bg3"], outline="", width=0)
        self.bar = canvas.create_rectangle(x, y, x, y+h, fill=color, outline="", width=0)
        self._tick()

    def _tick(self):
        elapsed = (time.time() - self.start) * 1000
        t = min(elapsed / self.duration, 1.0)
        t_e = 1 - (1-t)**3
        self.canvas.coords(self.bar, self.x, self.y,
                           self.x + int(self.w * self.pct * t_e), self.y + self.h)
        if t < 1.0:
            self.canvas.after(16, self._tick)


class ParticleBG:
    def __init__(self, canvas, count=40):
        import random
        self.canvas  = canvas
        self.running = True
        self.pts     = []
        W = int(canvas["width"]); H = int(canvas["height"])
        for _ in range(count):
            x   = random.uniform(0, W)
            y   = random.uniform(0, H)
            vx  = random.uniform(-0.3, 0.3)
            vy  = random.uniform(-0.2, 0.2)
            r   = random.uniform(1.0, 2.4)
            col = random.choice(["#58a6ff", "#3fb950", "#a371f7",
                                 "#39d353", "#d29922", "#f85149"])
            item = canvas.create_oval(x-r, y-r, x+r, y+r, fill=col, outline="")
            self.pts.append({"x":x,"y":y,"vx":vx,"vy":vy,"r":r,
                             "item":item,"W":W,"H":H})
        self._tick()

    def _tick(self):
        if not self.running:
            return
        for p in self.pts:
            p["x"] += p["vx"]; p["y"] += p["vy"]
            if p["x"] < 0:    p["x"] = p["W"]
            if p["x"] > p["W"]: p["x"] = 0
            if p["y"] < 0:    p["y"] = p["H"]
            if p["y"] > p["H"]: p["y"] = 0
            r = p["r"]
            self.canvas.coords(p["item"], p["x"]-r, p["y"]-r,
                               p["x"]+r, p["y"]+r)
        self.canvas.after(30, self._tick)

    def stop(self):
        self.running = False


class HoverBtn:
    def __init__(self, parent, text, command, bg, hbg,
                 fg="#ffffff", font=("Segoe UI", 10, "bold"),
                 padx=20, pady=8):
        import tkinter as tk
        self.widget = tk.Button(parent, text=text, command=command,
                                bg=bg, fg=fg, font=font, relief="flat",
                                bd=0, padx=padx, pady=pady,
                                activebackground=hbg, activeforeground=fg,
                                cursor="hand2")
        self.widget.bind("<Enter>",  lambda e: self.widget.config(bg=hbg))
        self.widget.bind("<Leave>",  lambda e: self.widget.config(bg=bg))

    def pack(self, **kw):  self.widget.pack(**kw)
    def grid(self, **kw):  self.widget.grid(**kw)

# ══════════════════════════════════════════════════════════════
#  7.  SPLASH SCREEN
# ══════════════════════════════════════════════════════════════

def show_splash(root):
    import tkinter as tk
    W, H = 500, 280
    splash = tk.Toplevel(root)
    splash.overrideredirect(True)
    sw = splash.winfo_screenwidth(); sh = splash.winfo_screenheight()
    splash.geometry(f"{W}x{H}+{(sw-W)//2}+{(sh-H)//2}")
    splash.configure(bg=C["bg"])

    # top accent border
    tk.Frame(splash, bg=C["accent"], height=3).place(x=0, y=0, width=W)
    # bottom accent border
    tk.Frame(splash, bg=C["accent"], height=3).place(x=0, y=H-3, width=W)

    # particle bg
    cv = tk.Canvas(splash, width=W, height=H, bg=C["bg"], highlightthickness=0)
    cv.place(x=0, y=0)
    pbg = ParticleBG(cv, count=25)

    # Title
    tk.Label(splash, text="MCQ ANALYZER",
             font=("Segoe UI", 26, "bold"),
             bg=C["bg"], fg=C["accent"]).place(relx=0.5, rely=0.30, anchor="center")
    tk.Label(splash, text="Professional Question Bank Inspector",
             font=("Segoe UI", 10),
             bg=C["bg"], fg=C["text2"]).place(relx=0.5, rely=0.50, anchor="center")

    # progress bar bg
    pb_bg = tk.Canvas(splash, width=320, height=4, bg=C["bg3"],
                      highlightthickness=0, bd=0)
    pb_bg.place(relx=0.5, rely=0.72, anchor="center")
    bar_id = pb_bg.create_rectangle(0, 0, 0, 4, fill=C["accent"], outline="")

    status = tk.Label(splash, text="Initializing...", font=("Segoe UI", 8),
                      bg=C["bg"], fg=C["text2"])
    status.place(relx=0.5, rely=0.84, anchor="center")

    msgs    = ["Loading modules...", "Preparing engine...", "Building UI...", "Ready!"]
    started = time.time()
    TOTAL   = 1500

    def _progress():
        e = (time.time()-started)*1000
        t = min(e/TOTAL, 1.0)
        t_e = 1-(1-t)**2
        pb_bg.coords(bar_id, 0, 0, int(320*t_e), 4)
        status.config(text=msgs[min(int(t*len(msgs)), len(msgs)-1)])
        if t < 1.0:
            splash.after(20, _progress)
        else:
            pbg.stop()
            splash.after(150, splash.destroy)

    splash.after(40, _progress)
    return splash

# ══════════════════════════════════════════════════════════════
#  8.  MAIN DASHBOARD WINDOW
# ══════════════════════════════════════════════════════════════

def show_dashboard(report_text, filepath, results, layout=None,
                   is_trial=False, lic_info=None):
    import tkinter as tk
    from tkinter import filedialog, messagebox

    root = tk.Tk()
    root.title("MCQ Analyzer  —  Professional Dashboard")
    root.geometry("1100x720")
    root.minsize(900, 600)
    root.configure(bg=C["bg"])
    try:
        root.state("zoomed")
    except Exception:
        pass

    splash = show_splash(root)
    root.wait_window(splash)

    # ── data shortcuts ────────────────────────────────────────
    q   = results["questions"];  tot = results["total"]
    mo  = results["missing_options"]; ma = results["missing_answer"]
    me  = results["missing_expln"];   mq = results["missing_qnums"]
    dg  = results["duplicates"]
    ivo = results.get("invalid_opts", [])
    iva = results.get("invalid_ans",  [])
    nums_p   = sorted(x["num"] for x in q)
    q_range  = f"Q{nums_p[0]} – Q{nums_p[-1]}" if nums_p else "N/A"
    t_issues = len(mo)+len(ma)+len(me)+len(mq)+len(dg)+len(ivo)+len(iva)
    health   = max(0, round(100 - t_issues/(max(tot,1))*30))

    # ── LEFT SIDEBAR ──────────────────────────────────────────
    sidebar = tk.Frame(root, bg=C["bg2"], width=230)
    sidebar.pack(side="left", fill="y")
    sidebar.pack_propagate(False)

    # logo
    logo = tk.Frame(sidebar, bg=C["bg2"])
    logo.pack(fill="x", pady=22)
    tk.Label(logo, text="MCQ", font=("Segoe UI", 22, "bold"),
             bg=C["bg2"], fg=C["accent"]).pack()
    tk.Label(logo, text="A N A L Y Z E R", font=("Segoe UI", 8, "bold"),
             bg=C["bg2"], fg=C["text2"]).pack()
    tk.Label(logo, text="Professional v2.0", font=("Segoe UI", 7),
             bg=C["bg2"], fg=C["text2"]).pack(pady=(2,0))

    tk.Frame(sidebar, bg=C["border"], height=1).pack(fill="x", padx=18, pady=2)

    # file chip
    fchip = tk.Frame(sidebar, bg=C["bg3"], pady=10)
    fchip.pack(fill="x", padx=14, pady=8)
    tk.Label(fchip, text="FILE", font=("Segoe UI", 7, "bold"),
             bg=C["bg3"], fg=C["text2"]).pack(anchor="w", padx=10)
    tk.Label(fchip, text=os.path.basename(filepath),
             font=("Segoe UI", 9, "bold"),
             bg=C["bg3"], fg=C["accent"], wraplength=195, justify="left").pack(anchor="w", padx=10)
    tk.Label(fchip, text=q_range,
             font=("Segoe UI", 8),
             bg=C["bg3"], fg=C["text2"]).pack(anchor="w", padx=10)

    tk.Frame(sidebar, bg=C["border"], height=1).pack(fill="x", padx=18, pady=2)

    # nav
    nav = tk.Frame(sidebar, bg=C["bg2"])
    nav.pack(fill="x", pady=6)

    active_tab = [None]
    tab_widgets = {}

    NAV_ITEMS = [
        ("dashboard", "Dashboard",        "▣"),
        ("options",   "Missing Options",  "◈"),
        ("answers",   "Missing Answers",  "◉"),
        ("expln",     "Missing Explns",   "◎"),
        ("missing",   "Missing Numbers",  "◌"),
        ("dupes",     "Duplicates",       "⊞"),
        ("invalid",   "Invalid Letters",  "⚠"),
        ("layout",    "Layout Check",     "⊡"),
        ("report",    "Full Report",      "≡"),
    ]

    for tab_id, label, icon in NAV_ITEMS:
        row = tk.Frame(nav, bg=C["bg2"], cursor="hand2")
        row.pack(fill="x", padx=10, pady=1)
        indicator = tk.Frame(row, bg=C["bg2"], width=3)
        indicator.place(x=0, y=0, height=38, width=3)
        lbl = tk.Label(row, text=f"  {icon}  {label}",
                       font=("Segoe UI", 10),
                       bg=C["bg2"], fg=C["text2"],
                       anchor="w", padx=10, pady=10)
        lbl.pack(fill="x")
        tab_widgets[tab_id] = (row, lbl, indicator)

        def _enter(e, r=row, l=lbl, t=tab_id):
            if active_tab[0] != t:
                r.config(bg=C["bg3"]); l.config(bg=C["bg3"])
        def _leave(e, r=row, l=lbl, t=tab_id):
            if active_tab[0] != t:
                r.config(bg=C["bg2"]); l.config(bg=C["bg2"])
        def _click(e, t=tab_id):
            switch_tab(t)

        for w in (row, lbl):
            w.bind("<Enter>",   _enter)
            w.bind("<Leave>",   _leave)
            w.bind("<Button-1>",_click)

    # health score
    tk.Frame(sidebar, bg=C["border"], height=1).pack(side="bottom", fill="x",
                                                      padx=18, pady=0)
    hf = tk.Frame(sidebar, bg=C["bg2"])
    hf.pack(side="bottom", fill="x", padx=16, pady=14)
    tk.Label(hf, text="FILE HEALTH", font=("Segoe UI", 7, "bold"),
             bg=C["bg2"], fg=C["text2"]).pack(anchor="w")
    hcol = C["teal"] if health >= 80 else C["orange"] if health >= 50 else C["red"]
    tk.Label(hf, text=f"{health}%",
             font=("Segoe UI", 30, "bold"),
             bg=C["bg2"], fg=hcol).pack(anchor="w")
    hbar_cv = tk.Canvas(hf, width=170, height=5, bg=C["bg3"],
                        highlightthickness=0, bd=0)
    hbar_cv.pack(anchor="w", pady=(3,0))
    hbar_cv.create_rectangle(0, 0, 0, 5, fill=hcol, outline="", tags="hb")
    def _hanim(t=0.0):
        t = min(t+0.035, 1.0)
        w = int(170 * (health/100) * (1-(1-t)**3))
        hbar_cv.coords("hb", 0, 0, w, 5)
        if t < 1.0:
            hbar_cv.after(16, lambda: _hanim(t+0.035))
    root.after(600, _hanim)

    # ── RIGHT CONTENT AREA ────────────────────────────────────
    main = tk.Frame(root, bg=C["bg"])
    main.pack(side="left", fill="both", expand=True)

    # header bar
    hdr = tk.Frame(main, bg=C["bg2"], height=56)
    hdr.pack(fill="x")
    hdr.pack_propagate(False)

    page_title_var = tk.StringVar(value="Dashboard")
    tk.Label(hdr, textvariable=page_title_var,
             font=("Segoe UI", 14, "bold"),
             bg=C["bg2"], fg=C["text"]).pack(side="left", padx=26, pady=16)

    hbtn = tk.Frame(hdr, bg=C["bg2"])
    hbtn.pack(side="right", padx=16, pady=10)

    def save_report():
        if is_trial and TRIAL_FEATURES.get("no_report_save"):
            messagebox.showwarning(
                "Trial Limitation",
                "Report saving is not available in Trial mode.\n\n"
                "Please activate your license to save reports.")
            return
        path = filedialog.asksaveasfilename(
            title="Save Report", defaultextension=".txt",
            filetypes=[("Text file","*.txt"),("All","*.*")],
            initialfile=os.path.splitext(os.path.basename(filepath))[0]+"_report.txt")
        if path:
            with open(path, "w", encoding="utf-8") as f:
                f.write(report_text)
            messagebox.showinfo("Saved", f"Report saved!\n{path}")

    def new_file():
        root.destroy(); main_entry()

    HoverBtn(hbtn, "  Save Report", save_report,
             C["btn_primary"], C["btn_prim_hov"]).pack(side="left", padx=4)
    HoverBtn(hbtn, "  New File", new_file,
             C["btn_second"], C["btn_sec_hov"]).pack(side="left", padx=4)
    HoverBtn(hbtn, "  Check Update", lambda: show_update_window(root),
             C["btn_second"], C["btn_sec_hov"]).pack(side="left", padx=4)

    # thin accent bar under header
    tk.Frame(main, bg=C["accent"], height=2).pack(fill="x")

    # ── TRIAL BANNER ──────────────────────────────────────────
    if is_trial:
        trial_bar = tk.Frame(main, bg="#2d2310", height=32)
        trial_bar.pack(fill="x")
        trial_bar.pack_propagate(False)
        tk.Label(trial_bar,
                 text=f"  TRIAL MODE  —  Only {TRIAL_FEATURES['max_questions']} questions analyzed  |  Layout check disabled  |  Report save disabled",
                 font=("Segoe UI", 8, "bold"),
                 bg="#2d2310", fg="#d29922").pack(side="left",
                                                   padx=12, pady=7)
        def _activate_now():
            show_license_window(root, on_success=lambda: (
                root.destroy(), main_entry()))
        tk.Button(trial_bar,
                  text="  Activate Now →",
                  command=_activate_now,
                  bg="#d29922", fg="#0d1117",
                  font=("Segoe UI", 8, "bold"),
                  relief="flat", bd=0, padx=12, pady=4,
                  cursor="hand2",
                  activebackground="#e3a11a").pack(side="right", padx=12, pady=5)

    # ── PAGE CONTAINER ────────────────────────────────────────
    pages_container = tk.Frame(main, bg=C["bg"])
    pages_container.pack(fill="both", expand=True)
    pages_dict = {}

    def scrollable_page(name):
        outer = tk.Frame(pages_container, bg=C["bg"])
        cv    = tk.Canvas(outer, bg=C["bg"], highlightthickness=0)
        sb    = tk.Scrollbar(outer, orient="vertical", command=cv.yview, bg=C["bg3"])
        inner = tk.Frame(cv, bg=C["bg"])
        win_id = cv.create_window((0,0), window=inner, anchor="nw")
        cv.configure(yscrollcommand=sb.set)
        inner.bind("<Configure>",
                   lambda e: cv.configure(scrollregion=cv.bbox("all")))
        cv.bind("<Configure>",
                lambda e: cv.itemconfig(win_id, width=e.width))
        cv.bind_all("<MouseWheel>",
                    lambda e: cv.yview_scroll(int(-1*(e.delta/120)), "units"))
        cv.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        pages_dict[name] = (outer, inner)
        return inner

    # ── PAGE: DASHBOARD ───────────────────────────────────────
    db = scrollable_page("dashboard")

    # stat cards grid
    cg = tk.Frame(db, bg=C["bg"])
    cg.pack(fill="x", padx=22, pady=(22,10))
    for i in range(3):
        cg.columnconfigure(i, weight=1)

    card_data = [
        ("TOTAL QUESTIONS", tot,      C["accent"],  C["card_blue"],   0, 0, q_range,               "Q"),
        ("MISSING OPTIONS", len(mo),  C["red"],     C["card_red"],    1, 0, "options < 4",          "!"),
        ("MISSING ANSWERS", len(ma),  C["orange"],  C["card_orange"], 2, 0, "no answer key",        "?"),
        ("MISSING EXPLN",   len(me),  C["purple"],  C["card_purple"], 0, 1, "no explanation",       "#"),
        ("MISSING Q NUMS",  len(mq),  C["teal"],    C["card_green"],  1, 1, "gaps in range",        "~"),
        ("DUPLICATES",      len(dg),  C["accent2"], C["card_blue"],   2, 1, "repeated questions",   "="),
        ("INVALID OPTIONS", len(ivo), C["red"],     C["card_red"],    0, 2, "non A/B/C/D letters",  "X"),
        ("INVALID ANSWERS", len(iva), C["orange"],  C["card_orange"], 1, 2, "non A/B/C/D answer",   "X"),
    ]

    for title, val, fg_col, bg_col, col, row, sub, ico in card_data:
        f = tk.Frame(cg, bg=bg_col,
                     highlightbackground=fg_col, highlightthickness=1)
        f.grid(row=row, column=col, padx=9, pady=9, sticky="nsew",
               ipadx=16, ipady=14)
        toprow = tk.Frame(f, bg=bg_col)
        toprow.pack(fill="x", padx=4, pady=(4,0))
        tk.Label(toprow, text=f"[{ico}]", font=("Segoe UI", 10, "bold"),
                 bg=bg_col, fg=fg_col).pack(side="left")
        tk.Label(toprow, text=title, font=("Segoe UI", 8),
                 bg=bg_col, fg=C["text2"]).pack(side="left", padx=6)
        val_lbl = tk.Label(f, text="0", font=("Segoe UI", 38, "bold"),
                           bg=bg_col, fg=fg_col)
        val_lbl.pack(anchor="w", padx=8)
        tk.Label(f, text=sub, font=("Segoe UI", 8),
                 bg=bg_col, fg=C["text2"]).pack(anchor="w", padx=8)
        root.after(350, lambda lbl=val_lbl, v=val, c=fg_col:
                   AnimatedCounter(lbl, v, c))

    # bar chart
    tk.Label(db, text="ISSUE BREAKDOWN", font=("Segoe UI", 9, "bold"),
             bg=C["bg"], fg=C["text2"]).pack(anchor="w", padx=30, pady=(16,4))

    chart_panel = tk.Frame(db, bg=C["bg2"],
                           highlightbackground=C["border"], highlightthickness=1)
    chart_panel.pack(fill="x", padx=20, pady=(0,20))

    bar_rows = [
        ("Missing Options",      len(mo),  C["red"]),
        ("Missing Answers",      len(ma),  C["orange"]),
        ("Missing Explanations", len(me),  C["purple"]),
        ("Missing Q Numbers",    len(mq),  C["teal"]),
        ("Duplicates",           len(dg),  C["accent2"]),
        ("Invalid Opt Letters",  len(ivo), C["red"]),
        ("Invalid Ans Letters",  len(iva), C["orange"]),
    ]
    max_v = max((v for _,v,_ in bar_rows), default=1) or 1

    bar_cv = tk.Canvas(chart_panel, bg=C["bg2"], highlightthickness=0, height=210)
    bar_cv.pack(fill="x", expand=True, padx=22, pady=18)

    def _draw_chart(event=None):
        bar_cv.delete("all")
        W = bar_cv.winfo_width() or 680
        bh = 24; gap = 10
        for i, (label, val, col) in enumerate(bar_rows):
            y  = 6 + i*(bh+gap)
            bar_cv.create_text(2, y+bh//2, text=label,
                               anchor="w", fill=C["text2"],
                               font=("Segoe UI", 9))
            bx = 210; bw = W - bx - 58
            bar_cv.create_rectangle(bx, y, bx+bw, y+bh,
                                    fill=C["bg3"], outline="", width=0)
            def _ab(c=col, v=val, x=bx, yy=y, w=bw, h=bh):
                AnimatedBar(bar_cv, x, yy, w, h, v/max_v, c, 950)
            root.after(550, _ab)
            bar_cv.create_text(W-2, y+bh//2, text=str(val),
                               anchor="e", fill=col,
                               font=("Segoe UI", 10, "bold"))

    bar_cv.bind("<Configure>", _draw_chart)
    root.after(300, _draw_chart)

    # ── ISSUE DETAIL PAGES ────────────────────────────────────
    def make_issue_page(name, title, icon, color, get_items):
        pg = scrollable_page(name)
        hf2 = tk.Frame(pg, bg=C["bg"])
        hf2.pack(fill="x", padx=24, pady=(24,6))
        tk.Label(hf2, text=f"{icon}  {title}",
                 font=("Segoe UI", 15, "bold"),
                 bg=C["bg"], fg=color).pack(side="left")

        items = get_items()
        if not items:
            tk.Frame(pg, bg=C["bg2"],
                     highlightbackground=C["border"],
                     highlightthickness=1).pack(fill="x", padx=20, pady=8,
                                                ipady=24, ipadx=20)
            ok_f = tk.Frame(pg, bg=C["bg2"])
            ok_f.pack_forget()
            # rebuild properly
            ok_card = tk.Frame(pg, bg=C["bg2"],
                               highlightbackground=C["accent2"],
                               highlightthickness=1)
            ok_card.pack(fill="x", padx=20, pady=8)
            tk.Label(ok_card,
                     text="  All Good!  No issues found in this category.",
                     font=("Segoe UI", 11),
                     bg=C["bg2"], fg=C["accent2"],
                     pady=20).pack(anchor="w", padx=16)
        else:
            tk.Label(pg, text=f"  {len(items)} issue(s) detected",
                     font=("Segoe UI", 9),
                     bg=C["bg"], fg=C["text2"]).pack(anchor="w", padx=24)
            for item in items:
                row = tk.Frame(pg, bg=C["bg2"],
                               highlightbackground=C["border"],
                               highlightthickness=1)
                row.pack(fill="x", padx=20, pady=4)
                badge = tk.Frame(row, bg=color, width=70)
                badge.pack(side="left", fill="y")
                badge.pack_propagate(False)
                tk.Label(badge, text=item["badge"],
                         font=("Segoe UI", 9, "bold"),
                         bg=color, fg=C["bg"]).pack(expand=True)
                tf = tk.Frame(row, bg=C["bg2"])
                tf.pack(side="left", fill="x", expand=True,
                        padx=14, pady=10)
                tk.Label(tf, text=item["main"],
                         font=("Segoe UI", 10, "bold"),
                         bg=C["bg2"], fg=C["text"],
                         anchor="w").pack(fill="x")
                if item.get("sub"):
                    tk.Label(tf, text=item["sub"],
                             font=("Segoe UI", 8),
                             bg=C["bg2"], fg=C["text2"],
                             anchor="w").pack(fill="x", pady=(2,0))

    def get_mo():
        return [{"badge": f"Q{n}",
                 "main":  f"Only {c} option(s)  (need 4)",
                 "sub":   next((x["statement"][:70] for x in q if x["num"]==n),"")}
                for n, c in mo]

    def get_ma():
        return [{"badge": f"Q{n}",
                 "main":  "Answer key is MISSING",
                 "sub":   next((x["statement"][:70] for x in q if x["num"]==n),"")}
                for n in ma]

    def get_me():
        return [{"badge": f"Q{n}",
                 "main":  "Explanation is MISSING",
                 "sub":   next((x["statement"][:70] for x in q if x["num"]==n),"")}
                for n in me]

    def get_mq():
        if not mq: return []
        ranges, s, e = [], mq[0], mq[0]
        for n in mq[1:]:
            if n == e+1: e = n
            else: ranges.append((s,e)); s = e = n
        ranges.append((s,e))
        return [{"badge": f"Q{a}-Q{b}" if a!=b else f"Q{a}",
                 "main":  f"{b-a+1} number(s) missing" if a!=b else "Number not present",
                 "sub":   ""}
                for a,b in ranges]

    def get_dg():
        return [{"badge": f"Grp {i+1}",
                 "main":  "Same statement: " + ", ".join(f"Q{n}" for n in g["qnums"]),
                 "sub":   g["preview"]}
                for i, g in enumerate(dg)]

    def get_ivo():
        return [{"badge": f"Q{n}",
                 "main":  f"Non-standard option letters: {', '.join(lts)}  (expected A/B/C/D)",
                 "sub":   next((x["statement"][:70] for x in q if x["num"]==n),"")}
                for n, lts in ivo]

    def get_iva():
        return [{"badge": f"Q{n}",
                 "main":  f"Answer '{lt}' is non-standard  (expected A/B/C/D)",
                 "sub":   next((x["statement"][:70] for x in q if x["num"]==n),"")}
                for n, lt in iva]

    make_issue_page("options", "Missing Options",      "◈", C["red"],    get_mo)
    make_issue_page("answers", "Missing Answer Keys",  "◉", C["orange"], get_ma)
    make_issue_page("expln",   "Missing Explanations", "◎", C["purple"], get_me)
    make_issue_page("missing", "Missing Q Numbers",    "◌", C["teal"],   get_mq)
    make_issue_page("dupes",   "Duplicate Questions",  "⊞", C["accent2"],get_dg)
    make_issue_page("invalid", "Invalid Letters",      "⚠", C["red"],
                   lambda: get_ivo() + get_iva())

    # ── LAYOUT CHECK PAGE ─────────────────────────────────────
    lp = scrollable_page("layout")
    tk.Label(lp, text="⊡  Document Layout Check",
             font=("Segoe UI", 15, "bold"),
             bg=C["bg"], fg=C["teal"]).pack(anchor="w", padx=24, pady=(24,8))

    if layout is None:
        no_lp = tk.Frame(lp, bg=C["bg2"],
                         highlightbackground=C["orange"], highlightthickness=1)
        no_lp.pack(fill="x", padx=20, pady=8)
        tk.Label(no_lp, text="  Layout analysis only available for .docx files.",
                 font=("Segoe UI", 11), bg=C["bg2"], fg=C["orange"],
                 pady=18).pack(anchor="w", padx=16)
    elif layout.get("error"):
        tk.Label(lp, text=f"  Error reading layout: {layout['error']}",
                 font=("Segoe UI", 10), bg=C["bg"], fg=C["red"]).pack(anchor="w", padx=24)
    else:
        mg = layout.get("margins", {})
        has_pg = layout.get("has_page_num", False)
        ls     = layout.get("line_spacing", "?")
        est_pg = layout.get("est_pages", "?")
        para_c = layout.get("para_count", "?")

        # animated margin diagram
        tk.Label(lp, text="PAGE MARGINS", font=("Segoe UI", 9, "bold"),
                 bg=C["bg"], fg=C["text2"]).pack(anchor="w", padx=28, pady=(8,4))

        margin_panel = tk.Frame(lp, bg=C["bg2"],
                                highlightbackground=C["border"], highlightthickness=1)
        margin_panel.pack(fill="x", padx=20, pady=(0,14))

        margin_cv = tk.Canvas(margin_panel, bg=C["bg2"], height=200,
                              highlightthickness=0)
        margin_cv.pack(fill="x", padx=20, pady=14)

        def _draw_margin_diagram(event=None):
            margin_cv.delete("all")
            W = margin_cv.winfo_width() or 600
            # page rect
            px, py, pw, ph = W//2-120, 20, 240, 160
            margin_cv.create_rectangle(px, py, px+pw, py+ph,
                                       fill=C["bg3"], outline=C["accent"], width=2)
            # content rect with animation
            def _anim_content(t=0.0):
                t = min(t+0.04, 1.0)
                te = 1-(1-t)**3
                pad = int(30*te)
                margin_cv.delete("content")
                margin_cv.create_rectangle(px+pad, py+pad,
                                           px+pw-pad, py+ph-pad,
                                           fill="#1c2d3f", outline=C["teal"],
                                           width=1, dash=(4,2), tags="content")
                if t < 1.0:
                    margin_cv.after(16, lambda: _anim_content(t+0.04))
            root.after(300, _anim_content)
            # labels
            for label, val, x, y in [
                (f"Top: {mg.get('top','?')} cm",    "", W//2, 12),
                (f"Bottom: {mg.get('bottom','?')} cm","", W//2, 185),
                (f"Left: {mg.get('left','?')} cm",  "", px-10, py+ph//2),
                (f"Right: {mg.get('right','?')} cm", "", px+pw+10, py+ph//2),
            ]:
                anch = "center" if "Top" in label or "Bottom" in label else \
                       ("e" if "Left" in label else "w")
                margin_cv.create_text(x, y, text=label, anchor=anch,
                                      fill=C["accent"], font=("Segoe UI", 9, "bold"))

        margin_cv.bind("<Configure>", _draw_margin_diagram)
        root.after(200, _draw_margin_diagram)

        # stat cards for layout
        layout_grid = tk.Frame(lp, bg=C["bg"])
        layout_grid.pack(fill="x", padx=20, pady=8)
        for i in range(3): layout_grid.columnconfigure(i, weight=1)

        lcard_data = [
            ("LINE SPACING",  str(ls),    C["accent"],  C["card_blue"],   0, 0, "per paragraph"),
            ("ESTIMATED PGS", str(est_pg),C["purple"],  C["card_purple"], 1, 0, "approximate"),
            ("PARAGRAPHS",    str(para_c),C["teal"],    C["card_green"],  2, 0, "total count"),
            ("PAGE NUMBERS",
             "Present" if has_pg else "MISSING",
             C["teal"] if has_pg else C["red"],
             C["card_green"] if has_pg else C["card_red"],
             0, 1, "in header/footer"),
            ("LEFT MARGIN",   f"{mg.get('left','?')} cm",  C["accent2"],C["card_blue"],   1, 1, ""),
            ("RIGHT MARGIN",  f"{mg.get('right','?')} cm", C["accent2"],C["card_blue"],   2, 1, ""),
        ]
        for title, val, fg_c, bg_c, col, row, sub in lcard_data:
            f2 = tk.Frame(layout_grid, bg=bg_c,
                          highlightbackground=fg_c, highlightthickness=1)
            f2.grid(row=row, column=col, padx=8, pady=8, sticky="nsew",
                    ipadx=14, ipady=12)
            tk.Label(f2, text=title, font=("Segoe UI", 8),
                     bg=bg_c, fg=C["text2"]).pack(anchor="w", padx=8)
            tk.Label(f2, text=val, font=("Segoe UI", 18, "bold"),
                     bg=bg_c, fg=fg_c).pack(anchor="w", padx=8)
            if sub:
                tk.Label(f2, text=sub, font=("Segoe UI", 7),
                         bg=bg_c, fg=C["text2"]).pack(anchor="w", padx=8)

        # Page number alert
        if not has_pg:
            alert = tk.Frame(lp, bg="#2d1c1c",
                             highlightbackground=C["red"], highlightthickness=1)
            alert.pack(fill="x", padx=20, pady=(4,12))
            tk.Label(alert,
                     text="  ⚠  Page numbers are MISSING from this document!",
                     font=("Segoe UI", 10, "bold"),
                     bg="#2d1c1c", fg=C["red"], pady=12).pack(anchor="w", padx=16)
        else:
            ok_f = tk.Frame(lp, bg=C["card_green"],
                            highlightbackground=C["teal"], highlightthickness=1)
            ok_f.pack(fill="x", padx=20, pady=(4,12))
            tk.Label(ok_f, text="  ✓  Page numbers detected in document.",
                     font=("Segoe UI", 10, "bold"),
                     bg=C["card_green"], fg=C["teal"], pady=12).pack(anchor="w", padx=16)

    # ── FULL REPORT PAGE ──────────────────────────────────────
    rp = scrollable_page("report")
    tk.Label(rp, text="≡  Full Text Report",
             font=("Segoe UI", 15, "bold"),
             bg=C["bg"], fg=C["accent"]).pack(anchor="w", padx=24, pady=(24,8))

    rp_card = tk.Frame(rp, bg=C["bg2"],
                       highlightbackground=C["border"], highlightthickness=1)
    rp_card.pack(fill="both", expand=True, padx=20, pady=(0,20))

    xsb = tk.Scrollbar(rp_card, orient="horizontal")
    xsb.pack(side="bottom", fill="x")
    txt = tk.Text(rp_card,
                  font=("Courier New", 9),
                  bg=C["bg"], fg=C["text"],
                  insertbackground=C["text"],
                  relief="flat", wrap="none",
                  xscrollcommand=xsb.set)
    txt.pack(fill="both", expand=True, padx=4, pady=4)
    xsb.config(command=txt.xview)

    txt.tag_configure("ok",   foreground=C["teal"])
    txt.tag_configure("warn", foreground=C["orange"])
    txt.tag_configure("err",  foreground=C["red"])
    txt.tag_configure("head", foreground=C["accent"],
                      font=("Courier New", 9, "bold"))
    txt.tag_configure("sep",  foreground=C["border"])

    for line in report_text.splitlines():
        if line.startswith("=") or line.startswith("-"):
            txt.insert("end", line+"\n", "sep")
        elif line.strip().isupper() and len(line.strip()) > 3:
            txt.insert("end", line+"\n", "head")
        elif "[OK]" in line:
            txt.insert("end", line+"\n", "ok")
        elif "MISSING" in line.upper() or "missing" in line:
            txt.insert("end", line+"\n", "warn")
        elif "Duplicate" in line or "Group" in line:
            txt.insert("end", line+"\n", "err")
        else:
            txt.insert("end", line+"\n")
    txt.config(state="disabled")

    # ── TAB SWITCH FUNCTION ───────────────────────────────────
    def switch_tab(name):
        page_title_var.set(name.replace("_", " ").title())
        for pname, (outer, inner) in pages_dict.items():
            outer.pack_forget()
        outer, inner = pages_dict[name]
        outer.pack(fill="both", expand=True)
        active_tab[0] = name
        for bname, (row, lbl, ind) in tab_widgets.items():
            if bname == name:
                row.config(bg=C["bg3"]); lbl.config(bg=C["bg3"], fg=C["accent"])
                ind.config(bg=C["accent"])
            else:
                row.config(bg=C["bg2"]); lbl.config(bg=C["bg2"], fg=C["text2"])
                ind.config(bg=C["bg2"])

    switch_tab("dashboard")
    root.mainloop()

# ══════════════════════════════════════════════════════════════
#  9.  CONSOLE HELPERS
# ══════════════════════════════════════════════════════════════

def cprint(text, color=None):
    CODES = {"red":"\033[91m","green":"\033[92m",
             "yellow":"\033[93m","cyan":"\033[96m","reset":"\033[0m"}
    if color and sys.stdout.isatty():
        print(f"{CODES.get(color,'')}{text}{CODES['reset']}")
    else:
        print(text)

# ══════════════════════════════════════════════════════════════
# 10.  PROCESSING WINDOW  (shown while analyzing)
# ══════════════════════════════════════════════════════════════

def show_processing_window(filepath):
    """Shows a professional processing/loading window while analyzing."""
    import tkinter as tk

    W, H = 520, 300
    proc = tk.Tk()
    proc.overrideredirect(True)
    proc.configure(bg="#0d1117")
    sw = proc.winfo_screenwidth(); sh = proc.winfo_screenheight()
    proc.geometry(f"{W}x{H}+{(sw-W)//2}+{(sh-H)//2}")
    proc.attributes("-topmost", True)

    # border
    tk.Frame(proc, bg="#58a6ff", height=3).place(x=0, y=0, width=W)
    tk.Frame(proc, bg="#238636", height=3).place(x=0, y=H-3, width=W)
    tk.Frame(proc, bg="#30363d", width=1).place(x=0, y=0, height=H)
    tk.Frame(proc, bg="#30363d", width=1).place(x=W-1, y=0, height=H)

    # particle canvas bg
    cv = tk.Canvas(proc, width=W, height=H, bg="#0d1117", highlightthickness=0)
    cv.place(x=0, y=0)
    pbg = ParticleBG(cv, count=20)

    # Dragon icon (drawn with canvas shapes)
    dc = tk.Canvas(proc, width=64, height=64, bg="#0d1117", highlightthickness=0)
    dc.place(relx=0.5, rely=0.18, anchor="center")
    # dragon body
    dc.create_oval(8, 20, 56, 56, fill="#161b22", outline="#58a6ff", width=2)
    dc.create_polygon(32,4, 20,22, 44,22, fill="#238636", outline="")
    dc.create_polygon(6,28, 2,18, 14,24, fill="#238636", outline="")
    dc.create_polygon(58,28, 62,18, 50,24, fill="#238636", outline="")
    dc.create_oval(18,30, 28,40, fill="#58a6ff", outline="")
    dc.create_oval(36,30, 46,40, fill="#58a6ff", outline="")
    dc.create_oval(21,33, 25,37, fill="#0d1117", outline="")
    dc.create_oval(39,33, 43,37, fill="#0d1117", outline="")
    dc.create_arc(22,42, 42,54, start=200, extent=140, style="arc",
                  outline="#f85149", width=2)

    tk.Label(proc, text="MCQ ANALYZER", font=("Segoe UI", 16, "bold"),
             bg="#0d1117", fg="#58a6ff").place(relx=0.5, rely=0.42, anchor="center")

    fname = os.path.basename(filepath)
    tk.Label(proc, text=f"Analyzing: {fname[:45]}{'...' if len(fname)>45 else ''}",
             font=("Segoe UI", 9), bg="#0d1117", fg="#8b949e").place(
             relx=0.5, rely=0.56, anchor="center")

    # status label
    status_lbl = tk.Label(proc, text="Reading file...",
                          font=("Segoe UI", 9, "bold"),
                          bg="#0d1117", fg="#3fb950")
    status_lbl.place(relx=0.5, rely=0.67, anchor="center")

    # progress bar
    pb_bg = tk.Canvas(proc, width=360, height=6, bg="#21262d",
                      highlightthickness=0, bd=0)
    pb_bg.place(relx=0.5, rely=0.78, anchor="center")
    bar_id = pb_bg.create_rectangle(0, 0, 0, 6, fill="#238636", outline="")

    steps_msgs = [
        (0.15, "Reading file..."),
        (0.35, "Parsing questions..."),
        (0.60, "Detecting issues..."),
        (0.80, "Building report..."),
        (1.00, "Done!"),
    ]
    _step_idx = [0]
    _started  = [time.time()]
    TOTAL_MS  = 1800

    def _tick_bar():
        elapsed = (time.time() - _started[0]) * 1000
        t = min(elapsed / TOTAL_MS, 1.0)
        t_e = 1 - (1-t)**2
        pb_bg.coords(bar_id, 0, 0, int(360*t_e), 6)
        # update status text
        for pct, msg in steps_msgs:
            if t <= pct:
                status_lbl.config(text=msg)
                break
        if t < 1.0:
            proc.after(20, _tick_bar)

    proc.after(30, _tick_bar)
    proc.update()
    return proc, pbg


# ══════════════════════════════════════════════════════════════
# 11b.  UPDATE CHECKER WINDOW
# ══════════════════════════════════════════════════════════════

def show_update_window(parent_win=None):
    """Professional update checker with AUTO DOWNLOAD."""
    import tkinter as tk
    import tempfile, subprocess

    W, H = 520, 380
    upd = tk.Toplevel(parent_win) if parent_win else tk.Tk()
    upd.overrideredirect(True)
    upd.configure(bg="#0d1117")
    sw = upd.winfo_screenwidth(); sh = upd.winfo_screenheight()
    upd.geometry(f"{W}x{H}+{(sw-W)//2}+{(sh-H)//2}")
    upd.attributes("-topmost", True)
    upd.grab_set()

    # ── borders ───────────────────────────────────────────────
    tk.Frame(upd, bg="#58a6ff", height=3).place(x=0, y=0, width=W)
    tk.Frame(upd, bg="#238636", height=3).place(x=0, y=H-3, width=W)
    tk.Frame(upd, bg="#30363d", width=1).place(x=0,   y=0, height=H)
    tk.Frame(upd, bg="#30363d", width=1).place(x=W-1, y=0, height=H)

    # ── title bar ─────────────────────────────────────────────
    tb = tk.Frame(upd, bg="#161b22", height=36)
    tb.place(x=1, y=3, width=W-2)
    tk.Label(tb, text="  MCQ Analyzer  —  Update Center",
             font=("Segoe UI", 10, "bold"),
             bg="#161b22", fg="#58a6ff").pack(side="left", padx=6, pady=8)
    close_x = tk.Button(tb, text="✕", command=upd.destroy,
              bg="#161b22", fg="#f85149",
              font=("Segoe UI", 10, "bold"),
              relief="flat", bd=0, padx=10,
              activebackground="#2d1c1c",
              cursor="hand2")
    close_x.pack(side="right")

    # ── animated icon ─────────────────────────────────────────
    ic = tk.Canvas(upd, width=56, height=56, bg="#0d1117", highlightthickness=0)
    ic.place(relx=0.5, rely=0.22, anchor="center")
    ic.create_oval(2, 2, 54, 54, fill="#161b22", outline="#58a6ff", width=2)
    ic_arrow = ic.create_text(28, 28, text="↓", font=("Segoe UI", 24, "bold"),
                              fill="#58a6ff")
    # pulse icon
    _ic_tick = [0]
    def _pulse_icon():
        _ic_tick[0] += 4
        colors = ["#58a6ff","#79c0ff","#a5d6ff","#79c0ff","#58a6ff"]
        c = colors[(_ic_tick[0]//8) % len(colors)]
        try: ic.itemconfig(ic_arrow, fill=c)
        except: pass
        upd.after(80, _pulse_icon)
    upd.after(100, _pulse_icon)

    # ── status label ──────────────────────────────────────────
    status_lbl = tk.Label(upd, text="Checking for updates...",
                          font=("Segoe UI", 13, "bold"),
                          bg="#0d1117", fg="#e6edf3")
    status_lbl.place(relx=0.5, rely=0.42, anchor="center")

    # ── version label ─────────────────────────────────────────
    ver_lbl = tk.Label(upd, text=f"Current version:  v{CURRENT_VERSION}",
                       font=("Segoe UI", 9),
                       bg="#0d1117", fg="#8b949e")
    ver_lbl.place(relx=0.5, rely=0.52, anchor="center")

    # ── notes label ───────────────────────────────────────────
    notes_lbl = tk.Label(upd, text="",
                         font=("Segoe UI", 9),
                         bg="#0d1117", fg="#8b949e",
                         wraplength=460, justify="center")
    notes_lbl.place(relx=0.5, rely=0.61, anchor="center")

    # ── progress bar ──────────────────────────────────────────
    pb_bg = tk.Canvas(upd, width=400, height=8, bg="#21262d",
                      highlightthickness=0)
    pb_bg.place(relx=0.5, rely=0.72, anchor="center")
    pb_bar = pb_bg.create_rectangle(0, 0, 0, 8, fill="#58a6ff", outline="")

    # ── progress percent label ────────────────────────────────
    pct_lbl = tk.Label(upd, text="",
                       font=("Segoe UI", 8),
                       bg="#0d1117", fg="#8b949e")
    pct_lbl.place(relx=0.5, rely=0.79, anchor="center")

    # ── spinning animation while checking ─────────────────────
    _spinning = [True]
    _spos = [0.0]
    def _spin():
        if not _spinning[0]: return
        _spos[0] = (_spos[0] + 0.035) % 1.2
        t = _spos[0]
        x1 = int(400 * max(0, t - 0.3))
        x2 = int(400 * min(1.0, t))
        pb_bg.coords(pb_bar, x1, 0, x2, 8)
        upd.after(16, _spin)
    upd.after(30, _spin)

    # ── close button ──────────────────────────────────────────
    close_btn = tk.Button(
        upd, text="  Close",
        command=upd.destroy,
        bg="#21262d", fg="#8b949e",
        font=("Segoe UI", 10),
        relief="flat", bd=0, padx=24, pady=8,
        cursor="hand2",
        activebackground="#30363d",
        activeforeground="#e6edf3"
    )
    close_btn.place(relx=0.5, rely=0.89, anchor="center")

    # ── AUTO DOWNLOAD FUNCTION ────────────────────────────────
    def _start_download(dl_url, latest_ver):
        """GitHub release se installer directly download karta hai."""

        # GitHub release URL se direct asset URL banana
        # dl_url = https://github.com/user/repo/releases/latest
        # Direct download URL format:
        # https://github.com/user/repo/releases/download/v2.1/MCQ_Analyzer_Setup_v2.1.exe
        asset_url = (f"https://github.com/{GITHUB_USER}/{GITHUB_REPO}"
                     f"/releases/download/v{latest_ver}/"
                     f"MCQ_Analyzer_Setup_v{latest_ver}.exe")

        save_path = os.path.join(
            os.path.expanduser("~"), "Downloads",
            f"MCQ_Analyzer_Setup_v{latest_ver}.exe")

        status_lbl.config(text="Downloading update...", fg="#d29922")
        pct_lbl.config(text="Starting download...")
        pb_bg.coords(pb_bar, 0, 0, 0, 8)
        pb_bg.itemconfig(pb_bar, fill="#d29922")
        close_btn.config(state="disabled", fg="#555")

        def _download():
            try:
                req = _urllib_req.Request(
                    asset_url,
                    headers={"User-Agent": "MCQ-Analyzer-Updater/2.0"})

                with _urllib_req.urlopen(req, timeout=30) as resp:
                    total = int(resp.headers.get("Content-Length", 0))
                    downloaded = 0
                    chunk_size = 8192

                    with open(save_path, "wb") as f:
                        while True:
                            chunk = resp.read(chunk_size)
                            if not chunk:
                                break
                            f.write(chunk)
                            downloaded += len(chunk)

                            if total > 0:
                                pct = downloaded / total
                                dl_mb  = downloaded / (1024*1024)
                                tot_mb = total / (1024*1024)

                                def _update_bar(p=pct, dm=dl_mb, tm=tot_mb):
                                    try:
                                        pb_bg.coords(pb_bar, 0, 0,
                                                     int(400*p), 8)
                                        pct_lbl.config(
                                            text=f"{dm:.1f} MB / {tm:.1f} MB  "
                                                 f"({int(p*100)}%)")
                                    except: pass
                                upd.after(0, _update_bar)

                # Download complete!
                def _done():
                    try:
                        pb_bg.coords(pb_bar, 0, 0, 400, 8)
                        pb_bg.itemconfig(pb_bar, fill="#3fb950")
                        status_lbl.config(
                            text="Download Complete!  Installing...",
                            fg="#3fb950")
                        pct_lbl.config(
                            text=f"Saved: {save_path}",
                            fg="#8b949e")
                        close_btn.config(state="normal", fg="#8b949e")
                        # Auto launch installer
                        upd.after(800, lambda: _launch_installer(save_path))
                    except: pass
                upd.after(0, _done)

            except Exception as e:
                def _err(err=str(e)):
                    try:
                        pb_bg.itemconfig(pb_bar, fill="#f85149")
                        status_lbl.config(
                            text="Download Failed!", fg="#f85149")
                        pct_lbl.config(text=f"Error: {err}", fg="#f85149")
                        ver_lbl.config(
                            text="Please download manually from GitHub.",
                            fg="#8b949e")
                        close_btn.config(state="normal", fg="#8b949e")
                        # fallback: open browser
                        webbrowser.open(dl_url)
                    except: pass
                upd.after(0, _err)

        threading.Thread(target=_download, daemon=True).start()

    def _launch_installer(path):
        """Installer launch karta hai aur app band karta hai."""
        try:
            status_lbl.config(text="Launching installer...", fg="#3fb950")
            upd.after(600, lambda: (
                subprocess.Popen([path], shell=True),
                upd.destroy(),
                sys.exit(0)      # purana app band
            ))
        except Exception as e:
            webbrowser.open(DOWNLOAD_URL)

    # ── CHECK UPDATE THREAD ───────────────────────────────────
    def _do_check():
        result = check_for_update()
        _spinning[0] = False

        def _show():
            pb_bg.coords(pb_bar, 0, 0, 400, 8)

            if result["error"]:
                # ── Connection error ──────────────────────────
                pb_bg.itemconfig(pb_bar, fill="#f85149")
                status_lbl.config(text="Connection Failed!", fg="#f85149")
                ver_lbl.config(
                    text=f"Could not reach update server.",
                    fg="#8b949e")
                notes_lbl.config(
                    text="Check your internet connection and try again.",
                    fg="#f85149")
                pct_lbl.config(text="")
                close_btn.place(relx=0.5, rely=0.89, anchor="center")

            elif result["has_update"]:
                # ── Update available — auto start download ────
                pb_bg.itemconfig(pb_bar, fill="#d29922")
                status_lbl.config(
                    text=f"Update Found!  v{result['latest']}",
                    fg="#d29922")
                ver_lbl.config(
                    text=f"v{result['current']}  →  v{result['latest']}",
                    fg="#e6edf3")
                notes_lbl.config(text=result["notes"], fg="#8b949e")
                pct_lbl.config(
                    text="Download starting automatically...",
                    fg="#d29922")

                # Auto start download after 1.5 seconds
                upd.after(1500, lambda: _start_download(
                    result["download_url"], result["latest"]))

            else:
                # ── Already up to date ────────────────────────
                pb_bg.itemconfig(pb_bar, fill="#3fb950")
                status_lbl.config(
                    text="You are up to date!",
                    fg="#3fb950")
                ver_lbl.config(
                    text=f"v{result['current']}  is the latest version.",
                    fg="#8b949e")
                notes_lbl.config(text="No updates available.", fg="#8b949e")
                pct_lbl.config(text="")
                close_btn.place(relx=0.5, rely=0.89, anchor="center")

        upd.after(0, _show)

    threading.Thread(target=_do_check, daemon=True).start()

    if parent_win is None:
        upd.mainloop()


# ══════════════════════════════════════════════════════════════
# 11c.  DRAGON LAUNCHER WINDOW
# ══════════════════════════════════════════════════════════════

def draw_dragon(canvas, x, y, scale=1.0, color_body="#161b22",
                color_accent="#58a6ff", color_fire="#f85149",
                color_wing="#238636"):
    """Draws a dragon using canvas shapes at position (x,y)."""
    s = scale
    # wings
    canvas.create_polygon(
        x, y+10*s,
        x-30*s, y-20*s,
        x-50*s, y+5*s,
        x-20*s, y+25*s,
        fill=color_wing, outline=color_accent, width=max(1,int(s)))
    canvas.create_polygon(
        x, y+10*s,
        x+30*s, y-20*s,
        x+50*s, y+5*s,
        x+20*s, y+25*s,
        fill=color_wing, outline=color_accent, width=max(1,int(s)))
    # tail
    canvas.create_line(x-8*s, y+40*s, x-25*s, y+65*s, x-15*s, y+75*s,
                       x-30*s, y+85*s,
                       fill=color_accent, width=max(2,int(2*s)), smooth=True)
    # body
    canvas.create_oval(x-28*s, y, x+28*s, y+55*s,
                       fill=color_body, outline=color_accent,
                       width=max(2,int(2*s)))
    # belly scales
    for i in range(3):
        canvas.create_oval(x-10*s, y+15*s+i*12*s,
                           x+10*s, y+25*s+i*12*s,
                           fill="#21262d", outline=color_accent,
                           width=max(1,int(s)))
    # neck
    canvas.create_oval(x-14*s, y-22*s, x+14*s, y+10*s,
                       fill=color_body, outline=color_accent,
                       width=max(2,int(2*s)))
    # head
    canvas.create_oval(x-20*s, y-55*s, x+20*s, y-18*s,
                       fill=color_body, outline=color_accent,
                       width=max(2,int(2*s)))
    # horns
    canvas.create_polygon(x-10*s, y-55*s, x-15*s, y-75*s, x-5*s, y-55*s,
                           fill=color_wing, outline="")
    canvas.create_polygon(x+10*s, y-55*s, x+15*s, y-75*s, x+5*s, y-55*s,
                           fill=color_wing, outline="")
    # eyes
    canvas.create_oval(x-12*s, y-48*s, x-4*s, y-40*s,
                       fill=color_accent, outline="")
    canvas.create_oval(x+4*s, y-48*s, x+12*s, y-40*s,
                       fill=color_accent, outline="")
    canvas.create_oval(x-10*s, y-46*s, x-6*s, y-42*s,
                       fill="#0d1117", outline="")
    canvas.create_oval(x+6*s, y-46*s, x+10*s, y-42*s,
                       fill="#0d1117", outline="")
    # fire breath
    canvas.create_polygon(
        x-5*s, y-22*s,
        x-20*s, y-30*s,
        x-35*s, y-20*s,
        x-28*s, y-15*s,
        x-40*s, y-8*s,
        x-20*s, y-12*s,
        fill=color_fire, outline="", stipple="")
    canvas.create_oval(x-38*s, y-32*s, x-22*s, y-18*s,
                       fill="#d29922", outline="")
    canvas.create_oval(x-35*s, y-30*s, x-24*s, y-20*s,
                       fill=color_fire, outline="")


def show_launcher():
    """Main launcher window with dragon logo, 2 buttons + license check."""
    import tkinter as tk
    from tkinter import filedialog, messagebox

    # ── Check license first ───────────────────────────────────
    is_licensed, is_trial, lic_info = check_license()

    W, H = 780, 520
    win  = tk.Tk()
    win.title("MCQ Analyzer  —  Professional Edition")
    win.geometry(f"{W}x{H}")
    win.resizable(False, False)
    win.configure(bg="#0d1117")
    sw = win.winfo_screenwidth(); sh = win.winfo_screenheight()
    win.geometry(f"{W}x{H}+{(sw-W)//2}+{(sh-H)//2}")

    # ── Custom title bar ──────────────────────────────────────
    title_bar = tk.Frame(win, bg="#161b22", height=38)
    title_bar.pack(fill="x", side="top")
    title_bar.pack_propagate(False)

    tk.Label(title_bar, text="  MCQ ANALYZER  —  Professional Edition",
             font=("Segoe UI", 10, "bold"),
             bg="#161b22", fg="#58a6ff").pack(side="left", padx=8)

    # close / minimise buttons
    def _close(): win.destroy(); sys.exit(0)
    def _mini():  win.iconify()

    tk.Button(title_bar, text="✕", command=_close,
              bg="#161b22", fg="#f85149",
              font=("Segoe UI", 11, "bold"),
              relief="flat", bd=0, padx=10, pady=4,
              activebackground="#2d1c1c",
              cursor="hand2").pack(side="right")
    tk.Button(title_bar, text="—", command=_mini,
              bg="#161b22", fg="#8b949e",
              font=("Segoe UI", 11),
              relief="flat", bd=0, padx=10, pady=4,
              activebackground="#21262d",
              cursor="hand2").pack(side="right")

    # drag title bar to move window
    _drag = {"x": 0, "y": 0}
    def _start_drag(e): _drag["x"] = e.x; _drag["y"] = e.y
    def _do_drag(e):
        dx = e.x - _drag["x"]; dy = e.y - _drag["y"]
        nx = win.winfo_x()+dx;  ny = win.winfo_y()+dy
        win.geometry(f"+{nx}+{ny}")
    title_bar.bind("<ButtonPress-1>",   _start_drag)
    title_bar.bind("<B1-Motion>",       _do_drag)

    # accent line under title
    tk.Frame(win, bg="#238636", height=2).pack(fill="x")

    # ── LEFT PANEL — Dragon + branding ───────────────────────
    left = tk.Frame(win, bg="#0d1117", width=340)
    left.pack(side="left", fill="y")
    left.pack_propagate(False)

    # particle bg on left panel
    lcv = tk.Canvas(left, width=340, height=H-42,
                    bg="#0d1117", highlightthickness=0)
    lcv.place(x=0, y=0)
    pbg = ParticleBG(lcv, count=28)

    # Dragon canvas
    dragon_cv = tk.Canvas(left, width=220, height=220,
                          bg="#0d1117", highlightthickness=0)
    dragon_cv.place(relx=0.5, rely=0.38, anchor="center")

    # animated dragon scale
    _dscale = [0.0]
    def _dragon_appear():
        _dscale[0] = min(_dscale[0]+0.05, 1.8)
        dragon_cv.delete("all")
        draw_dragon(dragon_cv, 110, 150, scale=_dscale[0])
        if _dscale[0] < 1.8:
            dragon_cv.after(16, _dragon_appear)
    win.after(200, _dragon_appear)

    # glow ring animation under dragon
    _glow = [0]
    def _pulse_glow():
        _glow[0] = (_glow[0] + 3) % 360
        r = 72 + int(8 * math.sin(math.radians(_glow[0])))
        try:
            lcv.delete("glow")
            lcv.create_oval(170-r, 268-r, 170+r, 268+r,
                            outline="#238636",
                            width=2, tags="glow")
        except Exception:
            pass
        lcv.after(30, _pulse_glow)
    win.after(400, _pulse_glow)

    # App name below dragon
    tk.Label(left, text="MCQ ANALYZER",
             font=("Segoe UI", 18, "bold"),
             bg="#0d1117", fg="#58a6ff").place(relx=0.5, rely=0.72, anchor="center")
    tk.Label(left, text="Professional Edition  v2.0",
             font=("Segoe UI", 9),
             bg="#0d1117", fg="#8b949e").place(relx=0.5, rely=0.80, anchor="center")
    tk.Label(left, text="Question Bank Inspector",
             font=("Segoe UI", 8),
             bg="#0d1117", fg="#30363d").place(relx=0.5, rely=0.87, anchor="center")

    # ── VERTICAL DIVIDER ─────────────────────────────────────
    tk.Frame(win, bg="#30363d", width=1).pack(side="left", fill="y")

    # ── RIGHT PANEL — controls ────────────────────────────────
    right = tk.Frame(win, bg="#0d1117")
    right.pack(side="left", fill="both", expand=True)

    # welcome text
    tk.Label(right, text="Welcome",
             font=("Segoe UI", 22, "bold"),
             bg="#0d1117", fg="#e6edf3").place(relx=0.5, rely=0.14, anchor="center")
    tk.Label(right,
             text="Select your MCQ file to begin\nthe analysis process.",
             font=("Segoe UI", 11),
             bg="#0d1117", fg="#8b949e",
             justify="center").place(relx=0.5, rely=0.26, anchor="center")

    # decorative line
    sep_cv = tk.Canvas(right, width=260, height=2, bg="#0d1117",
                       highlightthickness=0)
    sep_cv.place(relx=0.5, rely=0.37, anchor="center")
    sep_cv.create_rectangle(0,0,260,2, fill="#238636", outline="")

    # supported formats label
    tk.Label(right,
             text="Supported:   .docx   .doc   .txt",
             font=("Segoe UI", 9),
             bg="#0d1117", fg="#58a6ff").place(relx=0.5, rely=0.44, anchor="center")

    # ── OPEN FILE BUTTON ──────────────────────────────────────
    def open_file():
        fp = filedialog.askopenfilename(
            title="Select MCQ File — MCQ Analyzer",
            filetypes=[("Word / Text", "*.docx *.doc *.txt"),
                       ("Word Document","*.docx *.doc"),
                       ("Text File","*.txt"),
                       ("All Files","*.*")])
        if not fp:
            return

        # close launcher, show processing window
        win.withdraw()
        pbg.stop()

        proc_win, proc_pbg = show_processing_window(fp)

        def _run_analysis():
            try:
                text      = load_file(fp)
                questions = parse_mcqs(text)
                if not questions:
                    proc_win.after(0, lambda: (
                        proc_pbg.stop(),
                        proc_win.destroy(),
                        win.destroy(),
                        _fatal("No questions detected.\n\n"
                               "Ensure format:\n  Q 1. Statement\n"
                               "  (A) option...\n  Answer: A")
                    ))
                    return

                # ── Trial limit enforce ───────────────────────
                if is_trial and len(questions) > TRIAL_FEATURES["max_questions"]:
                    questions = questions[:TRIAL_FEATURES["max_questions"]]

                results = analyze(questions)

                # ── Trial: disable layout ─────────────────────
                layout = None
                if not is_trial:
                    layout = analyze_layout(fp)

                report  = build_report(results, fp, layout)

                # ── Trial watermark in report ─────────────────
                if is_trial:
                    report = (
                        "=" * 64 + "\n"
                        "  TRIAL VERSION — Max 20 questions analyzed\n"
                        "  Activate license to remove all limits.\n"
                        "=" * 64 + "\n\n"
                    ) + report

                time.sleep(0.3)

                def _open_dashboard():
                    proc_pbg.stop()
                    proc_win.destroy()
                    win.destroy()
                    show_dashboard(report, fp, results, layout,
                                   is_trial=is_trial,
                                   lic_info=lic_info)

                proc_win.after(400, _open_dashboard)
            except Exception as e:
                proc_win.after(0, lambda: (
                    proc_pbg.stop(),
                    proc_win.destroy(),
                    messagebox.showerror("Error", str(e)),
                    win.deiconify()
                ))

        t = threading.Thread(target=_run_analysis, daemon=True)
        t.start()

        # keep processing window alive
        proc_win.mainloop()

    # Open File button
    open_btn = tk.Button(
        right,
        text="  OPEN FILE",
        command=open_file,
        bg="#238636", fg="#ffffff",
        font=("Segoe UI", 13, "bold"),
        relief="flat", bd=0,
        padx=30, pady=14,
        cursor="hand2",
        activebackground="#2ea043",
        activeforeground="#ffffff",
        width=18
    )
    open_btn.place(relx=0.5, rely=0.57, anchor="center")
    open_btn.bind("<Enter>", lambda e: open_btn.config(bg="#2ea043"))
    open_btn.bind("<Leave>", lambda e: open_btn.config(bg="#238636"))

    # file type chips below button
    chips_frame = tk.Frame(right, bg="#0d1117")
    chips_frame.place(relx=0.5, rely=0.68, anchor="center")
    for fmt, col in [(".docx","#58a6ff"), (".doc","#a371f7"), (".txt","#3fb950")]:
        tk.Label(chips_frame, text=fmt,
                 font=("Segoe UI", 9, "bold"),
                 bg="#161b22", fg=col,
                 padx=10, pady=4,
                 relief="flat").pack(side="left", padx=4)

    # ── EXIT BUTTON ───────────────────────────────────────────
    exit_btn = tk.Button(
        right,
        text="  EXIT",
        command=_close,
        bg="#21262d", fg="#8b949e",
        font=("Segoe UI", 11),
        relief="flat", bd=0,
        padx=30, pady=10,
        cursor="hand2",
        activebackground="#30363d",
        activeforeground="#e6edf3",
        width=18
    )
    exit_btn.place(relx=0.5, rely=0.79, anchor="center")
    exit_btn.bind("<Enter>", lambda e: exit_btn.config(bg="#30363d", fg="#e6edf3"))
    exit_btn.bind("<Leave>", lambda e: exit_btn.config(bg="#21262d", fg="#8b949e"))

    # ── CHECK UPDATE BUTTON ───────────────────────────────────
    upd_btn = tk.Button(
        right,
        text="  Check for Updates",
        command=lambda: show_update_window(win),
        bg="#0d1117", fg="#58a6ff",
        font=("Segoe UI", 9),
        relief="flat", bd=0,
        padx=10, pady=4,
        cursor="hand2",
        activebackground="#161b22",
        activeforeground="#79c0ff"
    )
    upd_btn.place(relx=0.5, rely=0.88, anchor="center")
    upd_btn.bind("<Enter>", lambda e: upd_btn.config(fg="#79c0ff",
                 font=("Segoe UI", 9, "underline")))
    upd_btn.bind("<Leave>", lambda e: upd_btn.config(fg="#58a6ff",
                 font=("Segoe UI", 9)))

    # version footer
    tk.Label(right, text="MCQ Analyzer  v2.0  |  Professional Dashboard",
             font=("Segoe UI", 7),
             bg="#0d1117", fg="#30363d").place(relx=0.5, rely=0.94, anchor="center")

    # ── LICENSE STATUS BADGE ──────────────────────────────────
    if is_licensed:
        badge_text  = f"  LICENSED  —  {lic_info.get('plan','Pro').upper()}"
        badge_color = "#238636"
        badge_fg    = "#3fb950"
    else:
        badge_text  = f"  TRIAL  —  Max {TRIAL_FEATURES['max_questions']} Questions"
        badge_color = "#2d2310"
        badge_fg    = "#d29922"

    lic_badge = tk.Label(right,
                         text=badge_text,
                         font=("Segoe UI", 8, "bold"),
                         bg=badge_color, fg=badge_fg,
                         padx=10, pady=4)
    lic_badge.place(relx=0.5, rely=0.94, anchor="center")

    # ── ACTIVATE LICENSE BUTTON (only in trial) ───────────────
    if not is_licensed:
        def _open_activation():
            def _on_success():
                # Refresh launcher
                win.destroy()
                show_launcher()
            show_license_window(win, on_success=_on_success)

        act_lnk = tk.Button(
            right,
            text="  Activate License",
            command=_open_activation,
            bg="#0d1117", fg="#d29922",
            font=("Segoe UI", 8, "bold"),
            relief="flat", bd=0,
            padx=8, pady=3,
            cursor="hand2",
            activebackground="#161b22"
        )
        act_lnk.place(relx=0.5, rely=0.99, anchor="center")
        act_lnk.bind("<Enter>", lambda e: act_lnk.config(fg="#e3a11a"))
        act_lnk.bind("<Leave>", lambda e: act_lnk.config(fg="#d29922"))

    # bottom accent
    tk.Frame(win, bg="#58a6ff", height=2).pack(side="bottom", fill="x")

    win.mainloop()


# ══════════════════════════════════════════════════════════════
# 12.  ENTRY POINT
# ══════════════════════════════════════════════════════════════

def main_entry():
    if len(sys.argv) > 1 and os.path.isfile(sys.argv[1]):
        filepath  = sys.argv[1]
        cprint(f"\n  Loading  : {filepath}", "cyan")
        text      = load_file(filepath)
        questions = parse_mcqs(text)
        if not questions:
            _fatal("No questions detected.")
        results = analyze(questions)
        layout  = analyze_layout(filepath)
        report  = build_report(results, filepath, layout)
        show_dashboard(report, filepath, results, layout)
    else:
        show_launcher()


if __name__ == "__main__":
    main_entry()
